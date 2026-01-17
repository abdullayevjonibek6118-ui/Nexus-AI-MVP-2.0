"""
Resume Parser Service

Extracts text content from various resume file formats (PDF, DOCX, TXT).
"""

from typing import Optional
import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file."""
    try:
        from pypdf import PdfReader
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return f"[PDF parsing error: {str(e)}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_bytes)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting DOCX text: {e}")
        return f"[DOCX parsing error: {str(e)}]"


def parse_resume(filename: str, file_bytes: bytes) -> str:
    """
    Universal resume parser that detects file type and extracts text.
    
    Args:
        filename: Original filename with extension
        file_bytes: Raw file bytes
        
    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()
    
    # PDF files
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    
    # DOCX files
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    
    # DOC files (older format) - attempt DOCX parser
    elif filename_lower.endswith('.doc'):
        # Try DOCX parser first, fallback to text
        try:
            return extract_text_from_docx(file_bytes)
        except:
            # Fallback to UTF-8 decode
            try:
                return file_bytes.decode('utf-8', errors='ignore')
            except:
                return "[DOC file - unable to parse. Please convert to DOCX or PDF.]"
    
    # Text files (TXT, RTF treated as text)
    elif filename_lower.endswith(('.txt', '.rtf')):
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['cp1251', 'latin-1', 'cp1252']:
                try:
                    return file_bytes.decode(encoding)
                except:
                    continue
            return "[Text file - encoding not supported]"
    
    # Unknown format
    else:
        # Attempt UTF-8 decode as last resort
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except:
            return f"[Unsupported file format: {filename}. Please use PDF, DOCX, or TXT.]"
